#!/usr/bin/env python3
"""
PDF to DOCX Converter using PDF-MCP Server Tools

This script processes a PDF file using the PDF-MCP server tools directly and converts it to DOCX format.
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

# Add the src directory to Python path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from pdf_mcp_server.tools.text_extraction import ReadTextTool, ExtractMetadataTool

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PDFToDocxConverter:
    def __init__(self, pdf_path, output_dir):
        self.pdf_path = Path(pdf_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Initialize tools
        self.text_tool = ReadTextTool()
        self.metadata_tool = ExtractMetadataTool()
        
    def create_docx_from_text(self, text_content, metadata=None, output_path=None):
        """Create a DOCX document from extracted text"""
        if output_path is None:
            output_path = self.output_dir / f"{self.pdf_path.stem}.docx"
        
        # Create new document
        doc = Document()
        
        # Add title if available from metadata
        if metadata and 'title' in metadata and metadata['title']:
            title = doc.add_heading(metadata['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            title = doc.add_heading(f"Extracted from {self.pdf_path.name}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section if available
        if metadata:
            doc.add_heading('Document Information', level=1)
            info_table = doc.add_table(rows=0, cols=2)
            info_table.style = 'Table Grid'
            
            for key, value in metadata.items():
                if value and key != 'title':
                    row = info_table.add_row()
                    row.cells[0].text = key.replace('_', ' ').title()
                    row.cells[1].text = str(value)
        
        # Add main content
        doc.add_heading('Content', level=1)
        
        # Split text into paragraphs and add to document
        paragraphs = text_content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Check if it looks like a heading (short line, all caps, etc.)
                if len(para_text) < 100 and (para_text.isupper() or para_text.count(' ') < 5):
                    doc.add_heading(para_text, level=2)
                else:
                    doc.add_paragraph(para_text)
        
        # Save document
        doc.save(output_path)
        logger.info(f"DOCX document saved to: {output_path}")
        return output_path
    
    async def process_pdf(self):
        """Process PDF file and convert to DOCX"""
        logger.info(f"Processing PDF: {self.pdf_path}")
        
        if not self.pdf_path.exists():
            logger.error(f"PDF file not found: {self.pdf_path}")
            return None
        
        try:
            # Step 1: Extract text from PDF
            logger.info("Extracting text from PDF...")
            text_args = {
                "file_path": str(self.pdf_path),
                "method": "pymupdf",
                "include_metadata": True,
                "preserve_layout": True
            }
            
            text_result = await self.text_tool.execute(**text_args)
            
            if not text_result or 'content' not in text_result:
                logger.error("Failed to extract text from PDF")
                return None
            
            text_content = text_result['content']
            logger.info(f"Extracted {len(text_content)} characters of text")
            
            # Step 2: Extract metadata
            logger.info("Extracting metadata...")
            metadata_args = {
                "file_path": str(self.pdf_path)
            }
            
            try:
                metadata_result = await self.metadata_tool.execute(**metadata_args)
                metadata = metadata_result.get('metadata', {}) if metadata_result else {}
            except Exception as e:
                logger.warning(f"Failed to extract metadata: {e}")
                metadata = {}
            
            # Step 3: Create DOCX document
            logger.info("Creating DOCX document...")
            output_path = self.create_docx_from_text(text_content, metadata)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            import traceback
            traceback.print_exc()
            return None

async def main():
    """Main function"""
    # Input and output paths
    pdf_path = r"g:\E盘\工作项目文件\AI_Agent\Trae_Abroad\MCP_run_PDF\14531.pdf"
    output_dir = r"g:\E盘\工作项目文件\AI_Agent\Trae_Abroad\MCP_run_PDF"
    
    # Create converter and process PDF
    converter = PDFToDocxConverter(pdf_path, output_dir)
    result = await converter.process_pdf()
    
    if result:
        print(f"\n✅ Successfully converted PDF to DOCX: {result}")
    else:
        print("\n❌ Failed to convert PDF to DOCX")
        sys.exit(1)

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())