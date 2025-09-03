#!/usr/bin/env python3
"""
Simple PDF to DOCX Converter

This script directly uses PyMuPDF to extract text from PDF and convert to DOCX format.
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

try:
    import fitz  # PyMuPDF
except ImportError:
    print("Error: PyMuPDF not installed. Please install it with: pip install PyMuPDF")
    sys.exit(1)

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class SimplePDFToDocxConverter:
    def __init__(self, pdf_path, output_dir):
        self.pdf_path = Path(pdf_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def extract_text_from_pdf(self):
        """Extract text from PDF using PyMuPDF"""
        logger.info(f"Opening PDF: {self.pdf_path}")
        
        try:
            # Open PDF
            doc = fitz.open(str(self.pdf_path))
            
            # Extract metadata
            metadata = doc.metadata
            
            # Extract text from all pages
            all_text = ""
            page_texts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                page_texts.append({
                    'page_number': page_num + 1,
                    'text': text
                })
                all_text += f"\n\n--- Page {page_num + 1} ---\n\n{text}"
            
            doc.close()
            
            logger.info(f"Extracted text from {len(page_texts)} pages")
            logger.info(f"Total characters: {len(all_text)}")
            
            return {
                'text': all_text,
                'pages': page_texts,
                'metadata': metadata,
                'total_pages': len(page_texts)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return None
    
    def create_docx_from_text(self, extraction_result, output_path=None):
        """Create a DOCX document from extracted text"""
        if output_path is None:
            output_path = self.output_dir / f"{self.pdf_path.stem}.docx"
        
        # Create new document
        doc = Document()
        
        # Add title
        metadata = extraction_result.get('metadata', {})
        if metadata and metadata.get('title'):
            title = doc.add_heading(metadata['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            title = doc.add_heading(f"Extracted from {self.pdf_path.name}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section if available
        if metadata:
            doc.add_heading('Document Information', level=1)
            info_table = doc.add_table(rows=0, cols=2)
            info_table.style = 'Table Grid'
            
            for key, value in metadata.items():
                if value and key != 'title':
                    row = info_table.add_row()
                    row.cells[0].text = key.replace('_', ' ').title()
                    row.cells[1].text = str(value)
        
        # Add extraction summary
        doc.add_heading('Extraction Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Pages: {extraction_result['total_pages']}\n")
        summary_para.add_run(f"Total Characters: {len(extraction_result['text'])}\n")
        summary_para.add_run(f"Source File: {self.pdf_path.name}")
        
        # Add main content
        doc.add_heading('Content', level=1)
        
        # Process each page
        for page_info in extraction_result['pages']:
            page_num = page_info['page_number']
            text = page_info['text'].strip()
            
            if text:
                # Add page heading
                doc.add_heading(f'Page {page_num}', level=2)
                
                # Split text into paragraphs and add to document
                paragraphs = text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Check if it looks like a heading (short line, all caps, etc.)
                        if len(para_text) < 100 and (para_text.isupper() or para_text.count(' ') < 5):
                            doc.add_heading(para_text, level=3)
                        else:
                            doc.add_paragraph(para_text)
        
        # Save document
        doc.save(output_path)
        logger.info(f"DOCX document saved to: {output_path}")
        return output_path
    
    def process_pdf(self):
        """Process PDF file and convert to DOCX"""
        logger.info(f"Processing PDF: {self.pdf_path}")
        
        if not self.pdf_path.exists():
            logger.error(f"PDF file not found: {self.pdf_path}")
            return None
        
        try:
            # Step 1: Extract text from PDF
            logger.info("Extracting text from PDF...")
            extraction_result = self.extract_text_from_pdf()
            
            if not extraction_result:
                logger.error("Failed to extract text from PDF")
                return None
            
            # Step 2: Create DOCX document
            logger.info("Creating DOCX document...")
            output_path = self.create_docx_from_text(extraction_result)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            import traceback
            traceback.print_exc()
            return None

def main():
    """Main function"""
    # Input and output paths
    pdf_path = r"g:\E盘\工作项目文件\AI_Agent\Trae_Abroad\MCP_run_PDF\14531.pdf"
    output_dir = r"g:\E盘\工作项目文件\AI_Agent\Trae_Abroad\MCP_run_PDF"
    
    # Check if input file exists
    if not Path(pdf_path).exists():
        print(f"❌ Error: PDF file not found: {pdf_path}")
        sys.exit(1)
    
    # Create converter and process PDF
    converter = SimplePDFToDocxConverter(pdf_path, output_dir)
    result = converter.process_pdf()
    
    if result:
        print(f"\n✅ Successfully converted PDF to DOCX: {result}")
        print(f"📁 Output directory: {output_dir}")
    else:
        print("\n❌ Failed to convert PDF to DOCX")
        sys.exit(1)

if __name__ == "__main__":
    main()