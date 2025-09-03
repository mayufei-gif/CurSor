"""
DOCX Exporter for PDF-MCP pipeline results.

Builds a Word document (.docx) from the processing pipeline output using
python-docx. Supports paragraphs (text), tables (as Word tables), images,
and formulas (image + LaTeX caption fallback).

Public API:
    build_docx_from_pipeline(result: dict|BaseModel, out_path: str) -> str

Notes:
    - This module is intentionally tolerant of partial/missing fields.
    - It accepts both a dict and Pydantic model and normalizes to dict.
    - It stores the DOCX at out_path and returns the path.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

try:
    from pydantic import BaseModel  # type: ignore
except Exception:  # pragma: no cover
    BaseModel = object  # type: ignore


def _to_dict(obj: Any) -> Dict[str, Any]:
    if obj is None:
        return {}
    if isinstance(obj, dict):
        return obj
    if hasattr(obj, "model_dump"):
        return obj.model_dump()
    if isinstance(obj, BaseModel):  # type: ignore
        return obj.dict()
    # Last resort: best-effort
    return dict(obj)


def _add_paragraphs(doc: Document, text_result: Dict[str, Any]):
    # Prefer per-page paragraphs if present, otherwise fallback to full_text
    pages = text_result.get("pages") or []
    if pages:
        for page in pages:
            ptext = page.get("text") or ""
            for paragraph in ptext.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
                else:
                    doc.add_paragraph("")
            doc.add_paragraph("")
    else:
        full_text = text_result.get("full_text") or ""
        for paragraph in full_text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
            else:
                doc.add_paragraph("")


def _add_tables(doc: Document, tables_result: Dict[str, Any]):
    tables = tables_result.get("tables") or []
    if not tables:
        return
    for t in tables:
        data = t.get("data")
        if not data:
            # try csv_data
            csv_data = t.get("csv_data")
            if csv_data:
                rows = [row.split(",") for row in csv_data.strip().splitlines()]
                data = rows
        if not data:
            continue

        rows = len(data)
        cols = len(data[0]) if rows else 0
        if rows == 0 or cols == 0:
            continue

        headers = t.get("headers")
        has_header = bool(headers) and len(headers) == cols

        table = doc.add_table(rows=(rows + (1 if has_header else 0)), cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        row_offset = 0
        if has_header:
            hdr_cells = table.rows[0].cells
            for j, h in enumerate(headers):
                hdr_cells[j].text = str(h) if h is not None else ""
            row_offset = 1

        for i, r in enumerate(data):
            cells = table.rows[i + row_offset].cells
            # normalize row length
            if len(r) < cols:
                r = list(r) + [""] * (cols - len(r))
            for j, cell in enumerate(r):
                cells[j].text = str(cell) if cell is not None else ""
        doc.add_paragraph("")


def _add_images(doc: Document, images: List[Dict[str, Any]], caption_prefix: str):
    for img in images:
        path = img.get("image") or img.get("image_path")
        if not path:
            continue
        p = Path(path)
        if not p.exists():
            continue
        try:
            # Limit width to ~6 inches to fit page
            doc.add_picture(str(p), width=Inches(6))
        except Exception:
            continue
        caption = img.get("caption") or img.get("latex")
        if caption:
            doc.add_paragraph(f"{caption_prefix} {caption}")
        doc.add_paragraph("")


def build_docx_from_pipeline(result: Union[Dict[str, Any], Any], out_path: Union[str, Path]) -> str:
    """Build a .docx file from pipeline results.

    Args:
        result: dict or Pydantic model of ProcessingResponse or its content
        out_path: path to write the .docx

    Returns:
        The saved .docx path (string)
    """
    res = _to_dict(result)
    # Accept either the top-level ProcessingResponse or a dict with content already
    content = res.get("content") or res

    doc = Document()

    # Text
    text_result = content.get("text") or {}
    if text_result:
        _add_paragraphs(doc, text_result)

    # Tables
    tables_result = content.get("tables") or {}
    if tables_result:
        _add_tables(doc, tables_result)

    # Formulas (as images + caption)
    formulas_result = content.get("formulas") or {}
    formulas = formulas_result.get("formulas") or []
    if formulas:
        # Convert formulas into uniform image entries with optional latex
        imgs: List[Dict[str, Any]] = []
        for f in formulas:
            imgs.append({
                "image": f.get("image_path"),
                "latex": f.get("latex"),
            })
        _add_images(doc, imgs, caption_prefix="LaTeX:")

    # Images (optional future output)
    images = content.get("images") or []
    if images:
        _add_images(doc, images, caption_prefix="Figure:")

    out_path = str(out_path)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path

